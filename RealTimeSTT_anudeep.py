from RealtimeSTT import AudioToTextRecorder
from colorama import Fore, Back, Style
import colorama
import os
from docx import Document
#from REALTIMESTT_audio_recorder import AudioToTextRecorder  # Import your audio recorder class

class RealTimeSTT:
    def __init__(self):
        recorder_config = {
        'spinner': False,
        'model': 'large-v3',#'large-v2',
        'language': 'en',
        'silero_sensitivity': 0.4,
        'webrtc_sensitivity': 2,
        'post_speech_silence_duration': 0.4,
        'min_length_of_recording': 0,
        'min_gap_between_recordings': 0,
        'enable_realtime_transcription': True,
        'realtime_processing_pause': 0.2,
        'realtime_model_type': 'tiny.en',
        'on_realtime_transcription_update': self.on_realtime_transcription_update, 
        }
        self.recorder = AudioToTextRecorder(
            **recorder_config
        )
        self.transcribed_text = ""  # To store the accumulated transcribed text
        
        colorama.init()

        self.full_sentences = [] # To store complete sentences for display
        self.previous_text = ""
        displayed_text = ""

    def on_realtime_transcription_update(self, transcribed_text: str):
        """Callback function to handle real-time transcription updates."""
        #print("going in on realtime transcription update method")
        if transcribed_text.strip():  # Check if the new text is not empty
            self.transcribed_text = transcribed_text + " "  # Accumulate new text
            self.text_detected(transcribed_text)  # Display the current text

    def save_to_markdown(self, text: str, filename: str):
        """Saves the transcribed text to a Markdown file."""
        with open(filename, 'w') as md_file:
            md_file.write("# Transcription\n\n")  # Add a title
            md_file.write(text.strip())  # Write the transcribed text
        print(f"Transcription saved to Markdown file: {filename}")

    def save_to_docx(self, text: str, filename: str):
        """Saves the transcribed text to a DOCX file."""
        doc = Document()
        doc.add_heading('Transcription of Speech to Text', level=1)  # Add a title
        doc.add_paragraph(text.strip())  # Add the transcribed text
        doc.save(filename)
        print(f"Transcription saved to DOCX file: {filename}")

    def start_transcription(self):
        """Starts the real-time transcription process."""
        print("Starting real-time transcription...")
        self.recorder.start()  # Start the recording

    def stop_transcription(self):
        """Stops the transcription and saves the text."""
        print("Stopping transcription...")
        self.recorder.stop()  # Implement this method in your recorder class
        # Save the transcribed text to files
        self.save_to_markdown(self.transcribed_text, "Output/transcription.md")
        self.save_to_docx(self.transcribed_text, "Output/transcription.docx")
    
    def clear_console(self):
        os.system('clear' if os.name == 'posix' else 'cls')
    ###### From old code #########
    
    def text_detected(self, text):
        """Display the detected text in the console."""
        if text.strip():  # Only display non-empty text
            self.clear_console()
            print(Fore.YELLOW + text.strip() + Style.RESET_ALL, end="", flush=True)

    def process_text(self, text):
        """Process the incoming text for display."""
        self.text_detected(text)

if __name__ == "__main__":
    stt = RealTimeSTT()
    try:
        print(" Initializing RealtimeSTT test..... ")
        stt.start_transcription()
        input("Press Enter to stop transcription...")  # Wait for user input to stop
    finally:
        stt.stop_transcription()
